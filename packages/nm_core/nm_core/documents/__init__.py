"""Order-document processing: fetch PDF → store → extract text → AI summary."""
from __future__ import annotations

import io
import logging
import re
import uuid

from sqlalchemy import select
from sqlalchemy.orm import Session

from nm_core import ai, ecourts
from nm_core.db.models.case import CaseOrder
from nm_core.db.models.document import Document
from nm_core.storage import get_storage

logger = logging.getLogger("nm_core.documents")

_BAD_NAME_CHARS = re.compile(r'[<>:"|?*\x00-\x1f/\\]')
_WIN_RESERVED = re.compile(r"^(?:CON|PRN|AUX|NUL|COM\d|LPT\d)$", re.IGNORECASE)


class InvalidFilename(ValueError):
    """A user-supplied document name failed validation."""


def validate_filename(new_name: str, *, original: str | None = None) -> str:
    """Validate + normalize a user rename, preserving the original extension.

    Ported from the legacy files rename: strips path separators, caps at 200 chars,
    rejects illegal chars / Windows reserved names / leading-trailing dot or space.
    Raises InvalidFilename on a bad name."""
    name = (new_name or "").replace("/", "").replace("\\", "")[:200].strip()
    if not name:
        raise InvalidFilename("empty name")
    if _BAD_NAME_CHARS.search(name):
        raise InvalidFilename("name contains invalid characters")
    stem = name.rsplit(".", 1)[0] if "." in name else name
    if _WIN_RESERVED.match(stem):
        raise InvalidFilename("name is a reserved system name")
    if name.startswith(".") or name.endswith((".", " ")):
        raise InvalidFilename("name cannot start/end with a dot or space")
    # Preserve the original extension if the new name dropped it.
    if original and "." in original:
        ext = "." + original.rsplit(".", 1)[1]
        if not name.lower().endswith(ext.lower()):
            name += ext
    return name


def extract_text_docx(data: bytes) -> str:
    try:
        import docx

        return "\n".join(p.text for p in docx.Document(io.BytesIO(data)).paragraphs)
    except Exception:  # noqa: BLE001
        return ""


def extract_text_any(data: bytes, *, content_type: str | None, filename: str | None) -> str:
    name = (filename or "").lower()
    ct = (content_type or "").lower()
    if name.endswith(".docx") or "wordprocessingml" in ct:
        return extract_text_docx(data)
    return extract_text(data)  # PDF / plain text


class DocumentRepository:
    def __init__(self, session: Session) -> None:
        self.s = session

    def create_upload(
        self, *, account_id, created_by, title: str, filename: str,
        content_type: str | None, storage_key: str,
    ) -> Document:
        doc = Document(
            account_id=account_id, created_by=created_by, kind="upload", title=title,
            filename=filename, content_type=content_type, storage_key=storage_key,
            status="pending",
        )
        self.s.add(doc)
        self.s.flush()
        return doc

    def get(self, document_id) -> Document | None:
        return self.s.get(Document, document_id)

    def list_for_accounts(self, account_ids, *, limit: int = 200) -> list[Document]:
        if not account_ids:
            return []
        return list(
            self.s.execute(
                select(Document)
                .where(Document.account_id.in_(account_ids))
                .order_by(Document.updated_at.desc())
                .limit(limit)
            ).scalars()
        )


def process_upload(session: Session, *, document_id) -> Document | None:
    """Extract text + AI-summarize an uploaded document (PDF or DOCX). No OCR."""
    doc = session.get(Document, document_id)
    if doc is None:
        return None
    data = get_storage().get(doc.storage_key)
    text = extract_text_any(data, content_type=doc.content_type, filename=doc.filename)
    doc.extracted_text = text or None
    doc.summary = ai.summarize_order(text)
    doc.page_count = page_count(data)
    doc.status = "processed"
    session.flush()
    return doc


def extract_text(pdf: bytes) -> str:
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(pdf)) as doc:
            return "\n".join((page.extract_text() or "") for page in doc.pages)
    except Exception:  # noqa: BLE001 — scanned/invalid PDFs → no text (OCR is future work)
        return ""


def page_count(pdf: bytes) -> int | None:
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(pdf)) as doc:
            return len(doc.pages)
    except Exception:  # noqa: BLE001
        return None


def process_order(session: Session, *, order_id: uuid.UUID) -> CaseOrder | None:
    """Download, store, and AI-summarize one order. Idempotent-ish (re-stores)."""
    order = session.get(CaseOrder, order_id)
    if order is None or not order.order_url:
        return order
    pdf = ecourts.fetch_pdf(order.order_url)
    order.file_path = get_storage().put(f"orders/{order.case_id}/{order.order_id}.pdf", pdf)
    order.page_count = page_count(pdf)
    order.summary = ai.summarize_order(extract_text(pdf))
    if not order.descriptive_name:
        order.descriptive_name = f"Order dated {order.order_date}"
    session.flush()
    return order


def list_unprocessed(session: Session, *, limit: int = 50) -> list[CaseOrder]:
    return list(
        session.execute(
            select(CaseOrder).where(CaseOrder.file_path.is_(None)).limit(limit)
        ).scalars()
    )


def _try_process(session: Session, order_id: uuid.UUID) -> bool:
    """Process one order, swallowing transient fetch/parse failures.

    On failure file_path stays NULL, so the order is retried next sweep instead of
    being stranded with a bad summary.
    """
    try:
        process_order(session, order_id=order_id)
        return True
    except Exception:  # noqa: BLE001 — one bad/transient order must not abort the batch
        logger.warning("order processing failed for %s", order_id, exc_info=True)
        return False


def process_for_case(session: Session, *, case_id: uuid.UUID) -> int:
    rows = session.execute(
        select(CaseOrder).where(CaseOrder.case_id == case_id, CaseOrder.file_path.is_(None))
    ).scalars()
    return sum(_try_process(session, order.id) for order in list(rows))


def process_pending(session: Session, *, limit: int = 50) -> int:
    return sum(
        _try_process(session, order.id) for order in list_unprocessed(session, limit=limit)
    )
